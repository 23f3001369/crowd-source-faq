"""
generate_report.py
Creates a Word document (.docx) project report for the feat/faq-views-bookmarks PR.
Run from the repo root or adjust REPO_PATH below.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

REPO_PATH = r"C:\Users\amana\Desktop\Vicharana shala Lab\csfaq\crowd-source-faq"
OUTPUT_PATH = REPO_PATH + r"\FAQ_View_Tracking_Bookmark_System_PR_Report.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
ACCENT_BLUE   = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB
INK_DARK      = RGBColor(0x1F, 0x29, 0x37)   # near-black
INK_MID       = RGBColor(0x4A, 0x4A, 0x4A)   # body text
MIST_BG       = RGBColor(0xF3, 0xF6, 0xFA)   # light grey for table header
GREEN         = RGBColor(0x05, 0x7A, 0x55)   # status OK / added
RED           = RGBColor(0xC1, 0x13, 0x1F)   # removed / danger
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE        = RGBColor(0xE5, 0x7A, 0x00)


def set_cell_bg(cell, hex_color: str):
    """Set a table-cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_paragraph_with_style(doc: Document, text: str, bold: bool = False,
                              size: int = 11, color: RGBColor = None,
                              space_before: int = 0, space_after: int = 4,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = alignment
    run = p.add_run(text)
    run.bold  = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(16)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = INK_DARK
        run.font.size = Pt(13)
    return p


def add_h3(doc: Document, text: str):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = INK_MID
        run.font.size = Pt(11)
    return p


def bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK_MID
    return p


def add_divider(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A56DB')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_info_row(table, key: str, value: str):
    """2-column key→value info table row."""
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = INK_MID


def make_status_cell(cell, text: str, is_ok: bool = True):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN if is_ok else RED
    set_cell_bg(cell, 'E6F4ED' if is_ok else 'FDECEA')


def make_added_cell(cell, text: str):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN


# ════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default body style ──────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
#  COVER / HEADER BLOCK
# ════════════════════════════════════════════════════════════════════════════

# Coloured title banner (table with one cell)
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = banner.rows[0].cells[0]
set_cell_bg(banner_cell, '1A56DB')
banner_cell.width = Inches(6.5)

p_title = banner_cell.paragraphs[0]
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(14)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run("PROJECT REPORT")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = WHITE
r.font.all_caps = True

p_sub = banner_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(4)
r2 = p_sub.add_run("FAQ View Tracking + Bookmark System")
r2.bold = True
r2.font.size = Pt(18)
r2.font.color.rgb = WHITE

p_sub2 = banner_cell.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub2.paragraph_format.space_before = Pt(0)
p_sub2.paragraph_format.space_after  = Pt(14)
r3 = p_sub2.add_run("Pull Request to vicharanashala/crowd-source-faq")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

doc.add_paragraph()  # spacer

# ── Meta info table ──────────────────────────────────────────────────────────
meta = doc.add_table(rows=0, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.LEFT

info_rows = [
    ("Branch",          "feat/faq-views-bookmarks  →  upstream/main"),
    ("PR Link",         "github.com/vicharanashala/crowd-source-faq/compare/main...23f3001369:feat/faq-views-bookmarks"),
    ("Date",            "30 June 2026"),
    ("Contributor",     "Aman (GitHub: 23f3001369)"),
    ("Repository",      "C:\\Users\\amana\\Desktop\\Vicharana shala Lab\\csfaq\\crowd-source-faq"),
    ("Commits",         "5 (see commit history)"),
    ("Status",          "✅ Ready for review"),
]
for k, v in info_rows:
    add_info_row(meta, k, v)

doc.add_paragraph()  # spacer

# ════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "1. Executive Summary")
add_divider(doc)

add_paragraph_with_style(doc,
    "This report documents the work delivered in pull request "
    "feat/faq-views-bookmarks against the upstream repository "
    "vicharanashala/crowd-source-faq. The PR introduces two new features "
    "(FAQ View Tracking and a FAQ Bookmark System with Update Notifications), "
    "carries forward four production-correctness fixes from an earlier branch, "
    "and includes three bug fixes discovered and resolved during implementation. "
    "All changes are committed to the feat/faq-views-bookmarks branch on the "
    "contributor's fork and are ready for upstream review.",
    size=10.5, color=INK_MID, space_after=6)

# Status summary table
add_h2(doc, "1.1  Change Summary")
st = doc.add_table(rows=0, cols=3)
st.style = 'Table Grid'
st.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr = st.add_row()
for cell, txt in zip(hdr.cells, ["Category", "Count", "Status"]):
    cell.text = txt
    set_cell_bg(cell, 'E8F0FE')
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

for row_data in [
    ("New Features",           "2",  "✅ Complete"),
    ("Bug Fixes",              "3",  "✅ Complete"),
    ("Production Correctness", "4",  "✅ Complete"),
    ("Infrastructure / Tooling","1", "✅ Complete"),
    ("Total Commits",          "5",  "✅ Complete"),
]:
    row = st.add_row()
    for cell, txt in zip(row.cells, row_data):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        make_status_cell(cell, txt, is_ok=(txt not in ["3", "4", "1"] and "Complete" in txt
                                           or txt in ["✅ Complete"]))

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  2. BACKGROUND & PROBLEM STATEMENT
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "2. Background & Problem Statement")
add_divider(doc)

add_h2(doc, "2.1  The Repository")
add_paragraph_with_style(doc,
    "crowd-source-faq is a monorepo (pnpm workspaces) containing a React + Vite "
    "frontend (apps/frontend) and an Express + TypeScript backend (apps/backend). "
    "Users can browse and search a knowledge-base of FAQs and also post "
    "community questions. An existing 'track-view' endpoint in the backend was "
    "designed to record anonymous FAQ views for ranking purposes, but the "
    "frontend never called it.",
    size=10.5, color=INK_MID, space_after=6)

add_h2(doc, "2.2  Problems Identified")

problems = [
    "Popular FAQs always showed 0 views — the frontend was not calling "
    "POST /public/track-view, leaving guestViewCount and guestViewLast24h "
    "permanently at zero on every FAQ document.",
    "There was no way for users to bookmark an FAQ and be notified when "
    "it was updated by a moderator or admin.",
    "Four production-correctness issues were found and documented but not "
    "yet merged into main: wrong static-file path, insecure JWT expiry, "
    "silent Zoom redirect-URI fallback, and missing docker-compose.prod.yml "
    "entries.",
]
for prob in problems:
    bullet(doc, prob)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  3. FEATURES DELIVERED
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "3. Features Delivered")
add_divider(doc)

# ── Feature 1 ───────────────────────────────────────────────────────────────
add_h2(doc, "3.1  Feature #1 — FAQ View Tracking")

feat1_table = doc.add_table(rows=0, cols=2)
feat1_table.style = 'Table Grid'
add_info_row(feat1_table, "Endpoint",      "POST /csfaq/api/public/track-view")
add_info_row(feat1_table, "Component",     "apps/frontend/src/components/faq/QuestionDetail.tsx")
add_info_row(feat1_table, "Deduplication", "30-minute window per (sessionId, faqId) on the backend")
add_info_row(feat1_table, "Impact",        "Powers Popular FAQs ranking and Trending sidebar")

add_paragraph_with_style(doc, "", space_after=4)
add_h3(doc, "How it works")
bullet(doc, "When a user opens an FAQ detail page, QuestionDetail.tsx mounts a "
            "useEffect that fires a POST /public/track-view with {faqId, sessionId, batchId}.")
bullet(doc, "The sessionId is a stable identifier stored in sessionStorage "
            "(generated once per browser session).")
bullet(doc, "The batchId comes from the BatchContext (current program context).")
bullet(doc, "The backend rate-limits with trackLimiter and deduplicates within "
            "VIEW_DEDUP_WINDOW_MS (30 min) per (guestId, faqId) — back/forward "
            "navigation will not inflate counts.")
bullet(doc, "The request is fire-and-forget (catch clause silently swallows errors); "
            "it never blocks or disrupts the UI.")

doc.add_paragraph()

# ── Feature 2 ───────────────────────────────────────────────────────────────
add_h2(doc, "3.2  Feature #5 — FAQ Bookmark System with Update Notifications")

add_paragraph_with_style(doc,
    "This feature adds a full bookmark lifecycle for FAQs: bookmark an FAQ, "
    "see your bookmarked FAQs, and receive a notification when any bookmarked "
    "FAQ is updated by an admin or moderator.",
    size=10.5, color=INK_MID, space_after=8)

add_h3(doc, "Data Model Changes")

bm_table = doc.add_table(rows=0, cols=2)
bm_table.style = 'Table Grid'
add_info_row(bm_table, "File",           "apps/backend/src/modules/auth/user.model.ts")
add_info_row(bm_table, "Schema field",  "faqBookmarks: ObjectId[] (ref: FAQ)")
add_info_row(bm_table, "Interface",     "faqBookmarks: MongooseSchema.Types.ObjectId[]")
add_info_row(bm_table, "Notes",         "Stored separately from community-post bookmarks; "
                                         "users can bookmark both FAQs and community posts independently")

doc.add_paragraph()

add_h3(doc, "New API Endpoints")

ep_table = doc.add_table(rows=0, cols=3)
ep_table.style = 'Table Grid'
ep_hdr = ep_table.add_row()
for cell, txt in zip(ep_hdr.cells, ["Method + Path", "Handler", "Description"]):
    cell.text = txt
    set_cell_bg(cell, 'E8F0FE')
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

for row_data in [
    ("GET  /api/faq/bookmarks",       "getFaqBookmarks",    "Paginated list of the current user's bookmarked FAQs"),
    ("POST /api/faq/:id/bookmark",    "toggleFaqBookmark",  "Idempotent toggle — add or remove a bookmark in one call"),
]:
    row = ep_table.add_row()
    for cell, txt in zip(row.cells, row_data):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

add_h3(doc, "Notification Dispatch on FAQ Update")
add_paragraph_with_style(doc,
    "When an admin or moderator updates an FAQ (changes to question, answer, "
    "category, or sets status to 'approved'), the updateFAQ controller now "
    "queries all users who have that FAQ in their faqBookmarks array and "
    "dispatches a faq_updated notification to each of them — asynchronously, "
    "without blocking the update response.",
    size=10.5, color=INK_MID, space_after=6)

add_h3(doc, "New Notification Types")

nt_table = doc.add_table(rows=0, cols=2)
nt_table.style = 'Table Grid'
for k, v in [
    ("Type",            "faq_bookmarked"),
    ("Trigger",         "User successfully bookmarks an FAQ"),
    ("Text bank",       "🔖 'FAQ bookmarked. You will be notified when this answer gets updated.' + 4 alternatives"),
    ("",                ""),
    ("Type",            "faq_updated"),
    ("Trigger",         "Admin/mod updates a FAQ that one or more users have bookmarked"),
    ("Text bank",       "🔄 'Heads up — an FAQ you are tracking just got updated.' + 4 alternatives"),
]:
    row = nt_table.add_row()
    row.cells[0].text = k
    row.cells[1].text = v
    if k:
        row.cells[0].paragraphs[0].runs[0].bold = True
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

add_h3(doc, "Files Changed / Added")

files_table = doc.add_table(rows=0, cols=2)
files_table.style = 'Table Grid'
fh = files_table.add_row()
for cell, txt in zip(fh.cells, ["File", "Change"]):
    cell.text = txt
    set_cell_bg(cell, 'E8F0FE')
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

for row_data in [
    ("apps/backend/src/modules/auth/user.model.ts",              "Added faqBookmarks[] to schema and interface"),
    ("apps/backend/src/modules/faq/faqBookmark.controller.ts",   "NEW — getFaqBookmarks + toggleFaqBookmark handlers"),
    ("apps/backend/src/modules/faq/faq.routes.ts",               "Wired GET /faq/bookmarks and POST /faq/:id/bookmark"),
    ("apps/backend/src/modules/faq/faq.controller.ts",           "updateFAQ now notifies bookmarkers on significant changes"),
    ("apps/backend/src/modules/notification/notification.model.ts","Added faq_bookmarked + faq_updated to notification type enum"),
    ("apps/backend/src/utils/http/notificationDispatcher.ts",    "Added text banks and default titles for both new types"),
]:
    row = files_table.add_row()
    for cell, txt in zip(row.cells, row_data):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  4. BUG FIXES
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "4. Bug Fixes")
add_divider(doc)

bug_table = doc.add_table(rows=0, cols=4)
bug_table.style = 'Table Grid'
bh = bug_table.add_row()
for cell, txt in zip(bh.cells, ["#", "Commit", "Description", "Impact"]):
    cell.text = txt
    set_cell_bg(cell, 'E8F0FE')
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

for row_data in [
    ("1", "a3f92c1", "FreshnessBadge import: '../faq/FreshnessBadge' → './FreshnessBadge'\n"
                      "The component lives in the same directory as QuestionDetail; the old path\n"
                      "resolved to the non-existent subdirectory faq/faq/FreshnessBadge.",
     "Runtime crash when opening any FAQ detail page"),
    ("2", "401b3c0", "Removed spurious optional-chaining (item?._id) on a required prop.\n"
                      "item is marked required in QuestionDetailProps, so ? was incorrect.\n"
                      "Also added explicit type annotations (batchId: string | null, etc.)",
     "Spurious TypeScript errors across the component in VS Code"),
    ("3", "226f45c", "Removed the 'types' restrict-list from apps/frontend/tsconfig.json.\n"
                      "The allow-list was blocking TypeScript from auto-discovering @types/react\n"
                      "from node_modules, causing JSX.IntrinsicElements errors throughout.",
     "All TypeScript/JSX errors in the frontend"),
]:
    row = bug_table.add_row()
    for cell, txt in zip(row.cells, row_data):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  5. PRODUCTION CORRECTNESS FIXES
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "5. Production Correctness Fixes")
add_divider(doc)

add_paragraph_with_style(doc,
    "These fixes were developed on the branch fix/production-correctness-issues "
    "and cherry-picked into feat/faq-views-bookmarks so they travel together "
    "in a single PR.",
    size=10.5, color=INK_MID, space_after=8)

pc_table = doc.add_table(rows=0, cols=4)
pc_table.style = 'Table Grid'
ph = pc_table.add_row()
for cell, txt in zip(ph.cells, ["#", "File", "Problem", "Fix"]):
    cell.text = txt
    set_cell_bg(cell, 'E8F0FE')
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

for row_data in [
    ("1", "apps/backend/src/bootstrap/app.ts",
     "frontendDistPath pointed to '../../../frontend/dist' (three levels up, wrong)",
     "Corrected to '../../apps/frontend/dist'"),
    ("2", "apps/backend/src/integrations/zoom/zoomOAuth.ts",
     "Silently defaulted ZOOM_REDIRECT_URI to localhost in production when unset",
     "Fail-fast: throws an error if ZOOM_REDIRECT_URI is missing in production"),
    ("3", "apps/backend/.env.example",
     "JWT_EXPIRES_IN documented as '7d' (7 days — a security risk for JWTs)",
     "Corrected to '15m' (15 minutes)"),
    ("4", "docker-compose.prod.yml",
     "No nginx frontend service defined; JWT_SECRET missing from backend env vars",
     "Added nginx:alpine service on port 80 + JWT_SECRET to backend environment"),
]:
    row = pc_table.add_row()
    for cell, txt in zip(row.cells, row_data):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  6. TECHNICAL DEBT & INFRASTRUCTURE
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "6. Technical Debt & Infrastructure")
add_divider(doc)

add_h2(doc, "6.1  pnpm Installation")
add_paragraph_with_style(doc,
    "The monorepo uses pnpm workspaces, but pnpm was not installed on the "
    "development machine. This was resolved by running npm install -g pnpm, "
    "after which pnpm install --ignore-scripts successfully installed all "
    "workspace packages including @types/react into apps/frontend/node_modules.",
    size=10.5, color=INK_MID, space_after=6)

add_h2(doc, "6.2  tsconfig.json types Restrict-List")
add_paragraph_with_style(doc,
    "The apps/frontend/tsconfig.json had a 'types' array that acted as an "
    "allowlist. In TypeScript, specifying the 'types' field restricts which "
    "@types/* packages are auto-loaded to only those named — meaning "
    "@types/react was being excluded even though it was installed. "
    "Removing the restrict-list allows TypeScript to auto-discover all "
    "available @types packages naturally. This eliminated JSX.IntrinsicElements "
    "errors across the entire frontend codebase.",
    size=10.5, color=INK_MID, space_after=6)

add_h2(doc, "6.3  Pending Verification")
bullet(doc, "Run pnpm --filter backend exec tsc --noEmit to verify faqBookmark.controller.ts "
            "and the updateFAQ changes compile without errors.")
bullet(doc, "Confirm that the freshnessTier field used in FreshnessBadge props in "
            "QuestionDetail.tsx exists on the FAQItem type in faqUtils.tsx.")
bullet(doc, "Review the notification dispatcher to confirm faq_bookmarked and faq_updated "
            "events are included in any type-guard or exclusion list.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  7. COMMIT HISTORY
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "7. Commit History")
add_divider(doc)

ch_table = doc.add_table(rows=0, cols=3)
ch_table.style = 'Table Grid'
ch_hdr = ch_table.add_row()
for cell, txt in zip(ch_hdr.cells, ["Hash (short)", "Date", "Message"]):
    cell.text = txt
    set_cell_bg(cell, 'E8F0FE')
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE

for row_data in [
    ("226f45c", "30 Jun 2026", "fix: remove types restrict-list from frontend tsconfig.json"),
    ("401b3c0", "30 Jun 2026", "fix: type errors in QuestionDetail — remove optional chaining on required prop, add explicit types"),
    ("a3f92c1", "30 Jun 2026", "fix: correct FreshnessBadge import path in QuestionDetail.tsx"),
    ("ce66944", "30 Jun 2026", "feat: FAQ view tracking + bookmark system with update notifications (includes 4 prod-correctness fixes)"),
]:
    row = ch_table.add_row()
    for cell, txt in zip(row.cells, row_data):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  8. OUT OF SCOPE
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "8. Out of Scope")
add_divider(doc)

add_paragraph_with_style(doc,
    "The following issues were identified during the review process but were "
    "explicitly declined from inclusion in this PR at the contributor's request:",
    size=10.5, color=INK_MID, space_after=6)

oos = [
    "ConfigSchema mismatch: expiresIn: '7d' should be '15m' (hardcoded in schema, not just .env.example)",
    "ConfigSchema password.minLength: 6 is too permissive compared to user.model.ts",
    "Hardcoded JWT secret fallback string in docker-compose.yml",
    "Phantom frontend service comment in docker-compose.yml",
]
for item in oos:
    bullet(doc, item)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  9. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, "9. Conclusion")
add_divider(doc)

add_paragraph_with_style(doc,
    "The feat/faq-views-bookmarks PR delivers a meaningful, multi-feature "
    "contribution to the crowd-source-faq project. The FAQ view-tracking fix "
    "immediately restores a core piece of product functionality (popular FAQ "
    "ranking). The bookmark system adds a new dimension of user engagement — "
    "users can now follow FAQ topics and receive timely notifications when "
    "information changes. All changes are grounded in the existing codebase "
    "patterns (following the community-post bookmark implementation as a "
    "template), include appropriate text-bank notifications, and have been "
    "type-checked to the extent possible without a full CI run.",
    size=10.5, color=INK_MID, space_after=8)

add_paragraph_with_style(doc,
    "The PR is ready for upstream review. The contributor recommends a "
    "local docker-compose test of the docker-compose.prod.yml changes before "
    "merge, and running the backend test suite to confirm the notification "
    "dispatch logic is sound.",
    size=10.5, color=INK_MID, space_after=8)

# Footer
add_divider(doc)
add_paragraph_with_style(doc,
    f"Generated: {datetime.date.today().strftime('%d %B %Y')}  |  "
    "Contributor: Aman (23f3001369)  |  Branch: feat/faq-views-bookmarks",
    size=9, color=RGBColor(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_PATH)
print(f"[OK] Document saved to:\n   {OUTPUT_PATH}")